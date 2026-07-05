import io
from docx import Document
from app.parser.base import BaseParser
from app.core.exceptions import ParserError


class DocxParser(BaseParser):
    """Parser extracting plain text from Microsoft Word (.docx) documents."""

    def parse(self, content: bytes) -> str:
        try:
            # Wrap content bytes in a BytesIO stream
            docx_file = io.BytesIO(content)
            doc = Document(docx_file)
            
            text_elements = []
            
            # 1. Extract text from standard paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_elements.append(paragraph.text)
            
            # 2. Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_cells_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_cells_text:
                        text_elements.append(" | ".join(row_cells_text))
            
            return "\n\n".join(text_elements)
        except Exception as e:
            raise ParserError("DOCX", "Failed to extract text from DOCX document.", {"original_error": str(e)})
